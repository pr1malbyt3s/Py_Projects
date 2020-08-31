import docx
import re
import xlsxwriter

#Vishing	Cloud Console	Internal Systems

#Create output sheet.
workbook = xlsxwriter.Workbook('test.xlsx')
worksheet = workbook.add_worksheet('Findings')
worksheet.write('A1', 'title')
worksheet.write('B1', 'finding_type')
worksheet.write('C1', 'severity')
worksheet.write('D1', 'recommendation')
worksheet.write('E1', 'details')
worksheet.write('F1', 'source')
worksheet.write('G1', 'tools')
worksheet.write('H1', 'likelihood')
worksheet.write('I1', 'impact')
worksheet.write('J1', 'additional_guidance')

def xls_write(row, column, doc_part, ws):
    for finding in doc_part:
        ws.write(row, column, finding)
        row += 1
    return row

def doc_parse(doc_name, start_row):
    #Load the document for parsing.
    document = docx.Document(doc_name)
    #Initiate findings list. Findings are recognized as paragraph headers by parser.
    findings_titles = []
    #test_list = []
    #for para in document.paragraphs:
    #    test_list.append(para.text)
    #Iterate through paragraphs and add findings to list.
    for para in document.paragraphs:
        findings_titles.append((re.search(r"[^:]*$", para.text)).group().strip())
    #Remove empty strings from findings.
    while('' in findings_titles):
        findings_titles.remove('')
        #Remove the ratings from titles.
        for i, finding in enumerate(findings_titles):
            findings_titles[i] = re.sub(r"\(.*?\)", '', finding).strip()

    #Parse all the findings tables and add content to master list.
    findings_list = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    findings_list.append(para.text)

    #Build the recommendations list
    severities = []
    details = []
    impacts = []
    likelihoods = []
    sources = []
    tools = []
    recommendations = []
    add_guides = []
    for i in range(0, len(findings_list)):
        if(findings_list[i] == "Risk Rating:"):
            severities.append(findings_list[i+1])
        elif(findings_list[i] == "Recommendation:"):
            recommendations.append(findings_list[i+1])
        elif(findings_list[i] == "Source:"):
            sources.append(findings_list[i+1])
        elif(findings_list[i] == "Tools:"):
            tools.append(findings_list[i+1])
        elif(findings_list[i] == "Details"):
            j = i+1
            #for x in range(0, 10):
            #    print(findings_list[j+x])
            #    print("BREAK")
            detail = ''
        #    print(findings_list[j])
        #    print("break")
            while(findings_list[j] != "Additional Guidance"):
                detail += ' ' + findings_list[j]
                j += 1
            details.append(detail.strip())
        elif(findings_list[i] == "Risk"):
            likelihoods.append(findings_list[i+1])
            impacts.append(findings_list[i+2])
        elif(findings_list[i] == "Additional Guidance"):
            j = i+1
            guidance = ''
            while(findings_list[j] != "Evidence"):
                guidance += ' ' + findings_list[j]
                j += 1
            add_guides.append(guidance.strip())
        else:
            continue
    r = start_row
    xls_write(r, 0, findings_titles, worksheet)
    xls_write(r, 2, severities, worksheet)
    xls_write(r, 3, recommendations, worksheet)
    xls_write(r, 4, details, worksheet)
    xls_write(r, 5, sources, worksheet)
    xls_write(r, 6, tools, worksheet)
    xls_write(r, 7, likelihoods, worksheet)
    xls_write(r, 8, impacts, worksheet)
    end_row = xls_write(r, 9, add_guides, worksheet)
    comp_name = []
    for i in range(0, len(details)):
        comp_name.append(doc_name.split('.d')[0])
    xls_write(r, 1, comp_name, worksheet)
    #print(findings_list)
    #print(len(findings_titles))
    #print(len(details))
    #for title in test_list:
    #    print(title)
    #for detail in details:
    #    print(detail)
    #    print("BREAK")
    return end_row + 1
#doc_parse("Cloud Console.docx", 1)
#doc_parse("Malware.docx", doc_parse("Privileged Account.docx", doc_parse("Wireless.docx", 1)))
#doc_parse("Physical.docx", 1)
doc_parse("Network Capture.docx", doc_parse("Phishing.docx", doc_parse("External Vulnerability.docx", doc_parse("Internal Vulnerability.docx", doc_parse("Physical.docx", doc_parse("External.docx", doc_parse("Web App.docx", doc_parse("Internal.docx", 1))))))))
workbook.close()
